#coding=UTF-8

import sys
import Tkconstants, tkFileDialog
import glob
import os
import shutil
import PIL.Image
import time
import tkMessageBox
import re
from docx import Document
from docx.shared import Inches
from Tkinter import *


class Application:
	
	def __init__(self, master=None):
		self.fontePadrao = ("Arial", "10")
   
		self.titulo = Label(master, text="Selecione a pasta com os prints")
		self.titulo["font"] = ("Arial", "10", "bold")
		self.titulo.grid(row=0, column=0, columnspan=3)
   
		self.pastaLabel = Label(master,text="Pasta: ", font=self.fontePadrao)
		self.pastaLabel.grid(row=1,column=0, sticky=W)
   
		self.caminho = Entry(master)
		self.caminho["width"] = 40
		self.caminho["font"] = self.fontePadrao
		self.caminho.grid(row=1,column=1, sticky=W+E)
		self.caminho.configure(state='readonly')
		
		self.selecionar = Button(master)
		self.selecionar['text'] = 'Selecionar pasta'
		self.selecionar["font"] = ("Calibri", "8")
		self.selecionar["width"] = 12
		self.selecionar["command"] = self.selecionarPasta
		self.selecionar.grid(row=1, column=2)
   
		self.batedor = Button(master)
		self.batedor["text"] = "Bater meta"
		self.batedor["font"] = ("Calibri", "8")
		self.batedor["width"] = 12
		self.batedor["command"] = self.bateMeta
		self.batedor.grid(row=2,column=0, columnspan=3)
		self.batedor.configure(state='disable')
		
		self.logTitle = Label(master, text="Log: ")
		self.logTitle["font"] = self.fontePadrao
		self.logTitle.grid_forget()
		
		self.hideButton = Button(master)
		self.hideButton["text"] = "Ocultar log"
		self.hideButton["font"] = ("Calibri", "8")
		self.hideButton["width"] = 10
		self.hideButton["command"] = self.ocultar
		self.hideButton.grid_forget()	
		
		self.log = Text(master)
		self.log["font"] = self.fontePadrao
		self.log.configure(state='disabled')
		self.log.grid_forget()	
		
		# defining options for opening a directory
		self.dir_opt = options = {}
		options['initialdir'] = 'C:\\'
		options['mustexist'] = False
		options['parent'] = root
		options['title'] = 'This is a title'
		
	#oculta log
	def ocultar(self):
		if self.hideButton['text'] == 'Ocultar log':
			self.log.grid_forget()
			self.hideButton['text'] = 'Mostrar log'
		else:
			self.hideButton['text'] = 'Ocultar log'
			app.log.grid(row=4,column=0, columnspan=3,sticky=W+E)
			

	#Método batedor de meta
	def bateMeta(self):
		caminhoPasta = self.caminho.get()
		executar(caminhoPasta)
		
	#seleciona a pasta
	def selecionarPasta(self):
		#faz alguma coisa
		
		# get filename
		directoryName = tkFileDialog.askdirectory()
		self.caminho.configure(state='normal')
		self.caminho.delete(0,END)
		self.caminho.insert(0,directoryName)
		self.caminho.configure(state='readonly')
		self.batedor.configure(state='normal')
		
class Questao(object):
	def __init__(self, enunciado, pergunta, resposta, caminho2):
		self.caminho2 = caminho2
		self.enunciado = enunciado
		self.pergunta = pergunta
		self.resposta = resposta
		
		self.criarDocx()
		
	def nome(self):
		stringDeAnalise = ''
		

		if not self.pergunta:
			stringDeAnalise = self.enunciado[0]
		else:
			stringDeAnalise = self.pergunta[0]
		
		if stringDeAnalise.find('_') == -1:
			app.log.configure(state='normal') 
			app.log.insert('Erro de sintaxe em ' + stringDeAnalise)
			app.log.configure(state='disabled')
			
		else:
			return stringDeAnalise[0:stringDeAnalise.find('_')] 
				
	def listaParaDocx(self, lista, document):
		
		
		if not os.path.isdir(self.caminho2 + '/temp'):
			os.mkdir(self.caminho2 + '/temp')
		
		for i in range(0,len(lista)):
			im=PIL.Image.open(self.caminho2 + '/' + lista[i])
			widht = im.size[0]
			height = im.size[1]
			
			
			#compacta a imagem
			im = im.resize((int(round(widht/2)),int(round(height/2))),PIL.Image.ANTIALIAS)
			
			#salva

			im.save(self.caminho2 + '/temp/' +lista[i]+'.jpg',optimize=True,quality=80)
			
			imw = im.size[0]/1.5
			imdpp = imw/72
			
			if imdpp < 2:
				document.add_picture(self.caminho2 + '/temp/' +lista[i]+'.jpg', width=Inches(2))
			if imdpp > 6:
				document.add_picture(self.caminho2 + '/temp/' +lista[i]+'.jpg', width=Inches(5.5))
			else:
				document.add_picture(self.caminho2 + '/temp/' +lista[i]+'.jpg')
				
	def criarDocx(self):
		document = Document('default.docx')
		document.add_heading('Enunciado', 0)
		
		self.listaParaDocx(self.enunciado, document)
		
		if self.pergunta:
			self.listaParaDocx(self.pergunta, document)
		
		if self.resposta:
			document.add_heading('Resposta', 0)
			self.listaParaDocx(self.resposta, document)
		
		fileName = self.nome()
		
		app.log.configure(state='normal') 
		app.log.insert(END, 'foi criado ' + fileName + '.docx: \n\n')
		app.log.insert(END, str(len(self.enunciado)) + ' print(s) de enunciado \n')
		app.log.insert(END, str(len(self.pergunta)) + ' print(s) de letra \n')
		app.log.insert(END, str(len(self.resposta)) + ' print(s) de resposta \n\n')
			
		app.log.configure(state='disabled')
		
		document.save(self.caminho2 + '/' + fileName + '.docx')


def executar(directoryName):
	
	erro = False
	inicio = time.time()
	app.logTitle.grid(row=3,column=0, columnspan=2,sticky=W)
	app.hideButton.grid(row=3,column=2)
	app.log.grid(row=4,column=0, columnspan=3,sticky=W+E)
	listaDeCaminhos=[]
	listaDePrints=[]
	listaDeQuestoes=[]		
	
	for file in glob.glob(directoryName + "/*.png"):
		listaDeCaminhos.append(file)
		listaDePrints.append(file[len(directoryName)+1: len(file)])
		
	listaDePrints.sort()
	listaDeCaminhos.sort()
	
	listaDePrints = sorted_nicely(listaDePrints)
	listaDeCaminhos = sorted_nicely(listaDeCaminhos)
	
	for i in range(0,len(listaDeCaminhos)-1):
		print str(listaDeCaminhos[i]) + '\n'
	
	
	enunciado = []
	pergunta = []
	resposta = []
	
	if not listaDePrints:
		app.log.configure(state='normal')
		app.log.insert(END,'Epa, nenhuma imagem detectada :/')
		app.log.configure(state='disabled')
		erro = True

	for i in range(0,len(listaDePrints)):
		
		indexDoUnderline = listaDePrints[i].find('_')
		
		if ('_' not in listaDePrints[i]):
			app.log.configure(state='normal')
			app.log.insert(END,'Faltou "_" em ' + listaDePrints[i])
			app.log.configure(state='disabled')
			erro = True
			break
			
		if listaDePrints[i][indexDoUnderline+1] == 'E' or listaDePrints[i][indexDoUnderline+1] == 'e':
			#O print eh um enunciado
			#Se a lista enunciado não estiver vazia e as questões forem diferentes, limpa ela
			if enunciado and enunciado[0][0:enunciado[0].find('_')] != listaDePrints[i][0:indexDoUnderline]:
				enunciado = []
				pergunta = []
				resposta = []
					
			#adiciona esse print a lista de enunciados		
			enunciado.append(listaDePrints[i])
				
			#Se a proxima questao tem nome diferente, é porque ela acabou e cria um docx	
			if i == len(listaDePrints) -1 or (listaDePrints[i][0:indexDoUnderline] != listaDePrints[i+1][0:indexDoUnderline] and ('_L' not in listaDePrints[i+1] or '_l' not in listaDePrints[i+1] or '_I' not in listaDePrints[i+1] or '_i' not in listaDePrints[i+1])):
				Questao(enunciado, pergunta, resposta, directoryName)
				
		elif listaDePrints[i][indexDoUnderline+1] == 'L' or listaDePrints[i][indexDoUnderline+1] == 'l' or listaDePrints[i][indexDoUnderline+1] == 'I' or listaDePrints[i][indexDoUnderline+1] == 'i':
			#O print eh uma letra ou item
			#Se a lista perguntas não estiver vazia e as questões forem diferentes, limpa ela
			if pergunta and pergunta[0][0:pergunta[0].find('_')] != listaDePrints[i][0:indexDoUnderline]:
				pergunta = []
				resposta = []
					
			#adiciona esse print a lista de enunciados		
			pergunta.append(listaDePrints[i])
				
			#Se a proxima questao tem nome diferente, é porque ela acabou e cria um docx	
			if i == len(listaDePrints) -1 or listaDePrints[i][0:indexDoUnderline] != listaDePrints[i+1][0:indexDoUnderline]:
				if enunciado:
					Questao(enunciado, pergunta, resposta, directoryName)
				else:
					app.log.configure(state='normal')
					app.log.insert(END,listaDePrints[i] + ' sem enunciado')
					app.log.configure(state='disabled')				
					erro = True 
					break
					
		elif listaDePrints[i][indexDoUnderline+1] == 'R' or listaDePrints[i][indexDoUnderline+1] == 'r':
			#O print eh uma resposta
			#Se a lista respostas não estiver vazia e as questões forem diferentes, limpa ela
			if resposta and resposta[0][0:resposta[0].find('_')] != listaDePrints[i][0:indexDoUnderline]:
				resposta = []
					
			#adiciona esse print a lista de enunciados		
			resposta.append(listaDePrints[i])
				
			#Se a proxima questao tem nome diferente, é porque ela acabou e cria um docx	
			if i == len(listaDePrints) -1 or listaDePrints[i][0:indexDoUnderline] != listaDePrints[i+1][0:indexDoUnderline]:
				if enunciado:
					Questao(enunciado, pergunta, resposta, directoryName)
				else:
					app.log.configure(state='normal')
					app.log.insert(END,listaDePrints[i] + ' sem enunciado ou item')
					app.log.configure(state='disabled')				
					erro = True 	
					break
		else:
			app.log.configure(state='normal')
			app.log.insert(END,listaDePrints[i] + ' com erro de sintaxe')
			app.log.configure(state='disabled')				
			erro = True 
			break		

	if not erro:
		if not os.path.isdir(directoryName + '/prints'):
			os.mkdir(directoryName + '/prints')		
		for i in range(0,len(listaDePrints)):
			if not os.path.isfile(directoryName + '/prints/' + listaDePrints[i]):
				shutil.move(listaDeCaminhos[i], directoryName + '/prints')
	
	if os.path.isdir(directoryName + '/temp'):
		shutil.rmtree(directoryName + '/temp')


	fim = time.time()
	app.log.configure(state='normal')	
	app.log.insert(END, '\n\n Tempo de execução : ' + str(fim-inicio) + ' segundos')
	
	if not erro:
		app.log.insert(END, '\n\n Sucesso total *--*')
	else:
		app.log.insert(END, '\n\n Erro na execução. Verifique o log e os nomes dos prints.')
	
	app.log.configure(state='disabled')	
	
def sorted_nicely( l ):
    """ Sorts the given iterable in the way that is expected.
 
    Required arguments:
    l -- The iterable to be sorted.
 
    """
    convert = lambda text: int(text) if text.isdigit() else text
    alphanum_key = lambda key: [convert(c) for c in re.split('([0-9]+)', key)]
    return sorted(l, key = alphanum_key)


root = Tk()
root.wm_title('Batedor de meta - RA')
app = Application(root)
#root.iconbitmap('myicon.ico')

root.mainloop()



